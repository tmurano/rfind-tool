#!/usr/bin/env python3
"""Generate RFind Web Tool v3 design specification report as .docx"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Style setup ──
style = doc.styles["Normal"]
font = style.font
font.name = "Yu Gothic"
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.4

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Yu Gothic"
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "IBM Plex Mono"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


# ════════════════════════════════════════════════════════
# Title
# ════════════════════════════════════════════════════════
title = doc.add_heading("RFind Web Tool v3 — 設計仕様書", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("作成日: 2026-03-30", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# ════════════════════════════════════════════════════════
# 1. 設計と仕様書
# ════════════════════════════════════════════════════════
doc.add_heading("1. 設計と仕様書", level=1)

# 1.1 概要
doc.add_heading("1.1 ツール概要", level=2)
doc.add_paragraph(
    "RFind Web Tool v3 は、遺伝子発現量（FPKM/TPM）マトリックスと臨床情報（Demographics）を入力とし、"
    "参照インデックスとの方向性一致度を Running Fisher 法（4方向）で定量化するクライアントサイド Web アプリケーションである。"
    "すべての計算はブラウザ上で完結し、サーバーへのデータ送信は行わない。複数ユーザーが同時に使用しても互いに影響しない。"
)

doc.add_heading("1.1.1 技術スタック", level=3)
add_table(
    ["項目", "内容"],
    [
        ["構成", "単一 HTML ファイル（CSS + JavaScript インライン）"],
        ["外部ライブラリ", "SheetJS (xlsx-0.20.3): Excel 読み書き\n"
         "Plotly.js (v2): インタラクティブ可視化"],
        ["対応言語", "日本語 / English（i18n 切替）"],
        ["UI テーマ", "ダークテーマ（CSS カスタムプロパティ）"],
        ["計算方式", "全てクライアントサイド（ブラウザ内 JavaScript）"],
        ["ファイル配信", "静的ファイル（HTTP サーバー推奨）"],
    ],
)

# 1.2 処理フロー
doc.add_heading("1.2 処理フロー（4ステップ + 結果出力）", level=2)

doc.add_heading("Step 1: FPKM マトリックスのアップロード", level=3)
doc.add_paragraph("遺伝子 × サンプルの発現量マトリックスをアップロードする。")
add_table(
    ["検証項目", "内容", "エラー時の挙動"],
    [
        ["ファイル形式", "CSV, TSV, Excel (.xlsx/.xls)", "対応外形式はパースエラー"],
        ["最小構造", "ヘッダー行 + データ行1行以上、2列以上", "エラー表示（日英対応）"],
        ["1列目ヘッダー", "Gene列として認識可能な名前\n（Gene, Symbol, GeneID, GeneName,\nEnsembl, Entrez, Probe, Name, ID 等）", "認識不可ならエラー＋検出ヘッダー値を表示"],
        ["空サンプル名", "ヘッダー行に空のサンプル名がないこと", "エラー表示"],
        ["数値検証", "発現量の50%以上が非数値なら拒否", "エラー＋非数値セル数を表示"],
        ["遺伝子名正規化", "HGNC + Ortholog マッピング自動適用\n正規化数をプレビューバッジに表示", "マッピング未読込時はupperCase のみ"],
        ["発現量単位", "FPKM, TPM を想定（単位の検証は行わない）", "—"],
    ],
)

doc.add_heading("Step 2: Demographics ファイルのアップロード", level=3)
doc.add_paragraph('サンプル ID と診断群（Dx）を含むファイルをアップロード。Dx = "Control" が FC 算出の基準。')
add_table(
    ["検証項目", "内容", "エラー時の挙動"],
    [
        ["ID列", "ヘッダー名で自動検出（列順不問）\nID, SampleID, Sample_ID, SampleName,\nSubject, SubjectID, BrNum, Sample 等", "検出不可→エラー＋ヘッダー一覧表示"],
        ["Dx列", "ヘッダー名で自動検出（列順不問）\nDx, Diagnosis, Diag, Group, Condition 等", "検出不可→エラー＋ヘッダー一覧表示"],
        ["Age列", "自動検出（任意）\nAge, AgeDeath, Age_at_Death 等\n検出された場合、結果テーブルとscatter plotに使用", "未検出でも正常動作"],
        ["Control 必須", 'Dx列に "Control" が最低1件必要', "0件→エラー＋検出Dx値一覧表示"],
        ["ID照合", "FPKM のサンプル ID と突合\n不一致サンプルは自動除外＋除外リスト表示", "0件マッチ→エラー＋両側ID例表示"],
        ["マッチControl", "マッチしたサンプル中にControlが最低1件", "0件→エラー"],
        ["追加列", "ID・Dx以外の全列を保持\nExcel出力時に含まれる", "—"],
    ],
)

doc.add_heading("Step 3: 参照インデックスのアップロード", level=3)
doc.add_paragraph("Running Fisher の参照シグネチャ（b1）となるインデックスファイルを複数アップロード。")
add_table(
    ["検証項目", "内容", "エラー時の挙動"],
    [
        ["必須列", "Symbol列 + Fold Change列（列順不問）", "検出不可→エラー＋ヘッダー一覧＋許容列名例"],
        ["Symbol列", "Gene, Symbol, GeneName, GeneID,\nProbe, Name 等を許容", "—"],
        ["FC列", "Fold Change, FC, LogFC, Log2FC,\nFoldChange, Fold_Change 等を許容", "—"],
        ["P-Value/Rank列", "あっても無視（使用しない）", "—"],
        ["FC形式自動判定", "Ratio / Log2FC / SignedFC を自動検出", "判定結果をpillに [SignedFC] 等で表示"],
        ["遺伝子名正規化", "HGNC + Ortholog マッピング適用\n正規化数を表示", "—"],
        ["複数ファイル", "一括アップロード・D&D対応\n個別削除ボタンあり", "—"],
    ],
)

doc.add_heading("1.2.1 FC 形式の自動判定ロジック", level=3)
doc.add_paragraph("インデックスファイルの Fold Change 値の形式を自動判定し、すべて SignedFC に統一。")
add_table(
    ["形式", "値の範囲", "検出条件", "変換処理"],
    [
        ["Ratio（パターン1）", "0 ~ +∞", "全値が正", "FC<1 → -1/FC\nFC≥1 → そのまま"],
        ["Log2FC（パターン2）", "-∞ ~ +∞", "負値あり、|値|<1 が存在", "正 → 2^fc\n負 → -(2^|fc|)"],
        ["SignedFC（パターン3）", "-∞~-1, +1~+∞", "負値あり、(-1,+1)内の値なし", "変換なし"],
    ],
)

doc.add_heading("Step 4: 解析実行", level=3)
doc.add_paragraph("サンプル × インデックスごとに以下のパイプラインを実行。")

doc.add_paragraph("(a) Control 平均 FPKM の算出", style="List Number")
doc.add_paragraph("各遺伝子について、Dx = Control の全サンプルの FPKM 平均を算出。")

doc.add_paragraph("(b) Fold Change (FC) の計算", style="List Number")
doc.add_paragraph("FC = サンプル FPKM ÷ Control 平均 FPKM")

doc.add_paragraph("(c) FC のフィルタリングと符号変換", style="List Number")
add_table(
    ["条件", "処理"],
    [
        ["Control平均 = 0 または NaN", "NA（除外）"],
        ["|FC| < 1.2（閾値: THRESHOLD_FC = 1.2）", "NA（除外）"],
        ["FC ≤ 1", "-1/FC に変換（負の符号 = downregulated）"],
        ["FC > 1", "そのまま（正 = upregulated）"],
    ],
)

doc.add_paragraph("(d) Running Fisher スコアの算出", style="List Number")
doc.add_paragraph("b1（参照インデックス）と b2（サンプルの FC）で Running Fisher を実行。b2 の有効遺伝子数が 5 未満の場合はスコア = 0。")

doc.add_paragraph("(e) バッチ処理と中止機能", style="List Number")
doc.add_paragraph(
    "2サンプルずつ setTimeout でバッチ処理し、UI をブロックしない。"
    "プログレスバーで進捗表示。赤い「解析を中止」ボタンでキャンセル可能。"
    "中止時は結果を表示せず、UI をリセット。"
)

# 1.3 Running Fisher
doc.add_heading("1.3 Running Fisher アルゴリズム詳細", level=2)

doc.add_paragraph(
    "Illumina BaseSpace Correlation Engine の Running Fisher アルゴリズム（Tech Note 970-2014-007）に準拠。"
    "GSEA と類似のランクベース enrichment だが、統計的有意性を Fisher's exact test で評価する点が異なる。"
)

doc.add_heading("1.3.1 入力", level=3)
add_table(
    ["パラメータ", "説明"],
    [
        ["b1", "参照インデックスの遺伝子リスト [{Symbol, FC}]"],
        ["b2", "サンプルの FC 変換済み遺伝子リスト [{Symbol, FC}]"],
        ["P", "遺伝子ユニバースサイズ = FPKM ∩ Index の共通遺伝子数\n（Illumina定義: P1∩P2）\n※暫定的に DEFAULT_P = 17,317 を使用（後日 commonGenes.length に変更予定）"],
        ["CLIP_LIMIT", "スコアクリッピング上限 = 300"],
    ],
)

doc.add_heading("1.3.2 アルゴリズムの流れ", level=3)

steps = [
    "b1, b2 それぞれを重複除去（Symbol の lowercase で一意化）し、FC=0 を除外",
    "各セットを UP (FC>0) と DN (FC<0) に分割",
    "4方向の組み合わせで fisherBidirectional を実行: UU, UD, DU, DD",
    "fisherBidirectional の内部処理（方向 dir1, dir2 について）:\n"
    "  a. b1 から dir1 方向の遺伝子を FC 順にランク付け（ranked1）\n"
    "  b. b2 から dir2 方向の遺伝子を FC 順にランク付け（ranked2）\n"
    "  c. Direction 1: ranked1 を上位からスキャン、ranked2 との累積重複を数える\n"
    "     → hit があった時のみ phyperUpper(cum, P, K, n) を計算（高速化最適化）\n"
    "     → 全ランク中の最小 p 値 (p1) と最適ランク (m1) を記録\n"
    "  d. Direction 2: ranked2 を逆方向にスキャンして同様に p2, m2 を取得\n"
    "  e. 双方向の平均: meanLogP = (-log10(p1) + -log10(p2)) / 2",
    "各方向のスコア = -log10(p) × sign（concordant: +1, discordant: -1）\n"
    "  各方向をクリッピング（±CLIP_LIMIT）",
    "Total = 4方向スコアの合計 → さらにクリッピング（±CLIP_LIMIT）",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}", style="List Number")

doc.add_heading("1.3.3 方向と符号", level=3)
add_table(
    ["方向", "b1", "b2", "符号", "意味"],
    [
        ["UU", "UP", "UP", "+1", "Concordant（方向一致）"],
        ["UD", "UP", "DOWN", "-1", "Discordant"],
        ["DU", "DOWN", "UP", "-1", "Discordant"],
        ["DD", "DOWN", "DOWN", "+1", "Concordant（方向一致）"],
    ],
)

doc.add_heading("1.3.4 phyperUpper の高速化", level=3)
doc.add_paragraph(
    "Running Fisher のスキャンでは、各ランクで phyperUpper（超幾何分布の上側確率）を計算する。"
    "元の実装では全ランクで計算していたが、以下の数学的性質を利用して高速化："
)
add_bullet("hit なし（overlap 不変）の場合: n が増加、overlap 同一 → p 値は必ず悪化 → スキップ可能", "")
add_bullet("hit あり（overlap 増加）の場合のみ: p 値が改善する可能性 → 計算実行", "")
doc.add_paragraph(
    "この最適化により phyperUpper の呼び出し回数が「全遺伝子数（数千回）」から"
    "「実際の overlap 数（数十回）」に激減し、263サンプル×1インデックスで 90秒→20秒 に高速化。"
    "結果は数学的に同一。"
)

# 1.4 HGNC + Ortholog
doc.add_heading("1.4 遺伝子名正規化", level=2)

doc.add_heading("1.4.1 HGNC Alias → Official Symbol", level=3)
doc.add_paragraph(
    "HGNC (HUGO Gene Nomenclature Committee) の公式データから、"
    "alias/旧名/ENSEMBL ID → 最新 approved symbol へのマッピングテーブルを使用。"
)
add_table(
    ["項目", "内容"],
    [
        ["ファイル", "hgnc_map.json（2.1 MB）"],
        ["エントリ数", "97,737"],
        ["対象", "Previous symbols, Alias symbols, Ensembl Gene ID → Approved Symbol"],
        ["Ambiguous alias", "1,654件を除外（1つの alias が複数の approved symbol に対応するもの）"],
        ["生成スクリプト", "scripts/build_hgnc_map.py\nHGNC カスタムダウンロードから TSV 取得→JSON 変換"],
        ["読込方式", "ページ読込時に非同期 fetch\nfile:// プロトコルでは読込不可（HTTP サーバー必須）"],
    ],
)

doc.add_heading("1.4.2 Ortholog マッピング（Mouse/Rat → Human）", level=3)
doc.add_paragraph(
    "RGD (Rat Genome Database) の Ensembl ortholog データから、"
    "マウス/ラットの遺伝子名 → ヒトの遺伝子名への変換テーブルを使用。"
    "大半の ortholog は uppercase 化で一致するため、symbol が実際に異なるもののみ格納。"
)
add_table(
    ["項目", "内容"],
    [
        ["ファイル", "ortholog_map.json（42 KB）"],
        ["エントリ数", "2,042（uppercase で異なる ortholog のみ）"],
        ["データソース", "RGD_ORTHOLOGS_Ensembl.txt\nMouse + Rat → Human を1ファイルで取得"],
        ["Ambiguous", "11件を除外"],
        ["生成スクリプト", "scripts/build_ortholog_map.py"],
    ],
)

doc.add_heading("1.4.3 normalizeGene() の処理順序", level=3)
doc.add_paragraph("FPKM およびインデックスの読込時に、各遺伝子名に対して以下の順序で変換:")
doc.add_paragraph("1. 入力を uppercase 化", style="List Number")
doc.add_paragraph("2. Ortholog マッピングで検索（Mouse/Rat → Human）→ ヒットすれば変換", style="List Number")
doc.add_paragraph("3. HGNC マッピングで検索（alias → official symbol）→ ヒットすれば変換", style="List Number")
doc.add_paragraph("4. どちらにもヒットしなければ uppercase のまま返す", style="List Number")

# 1.5 パラメータ
doc.add_heading("1.5 パラメータ一覧", level=2)
add_table(
    ["パラメータ", "値", "定義箇所", "備考"],
    [
        ["THRESHOLD_FC", "1.2", "computeSampleFC()", "|FC| < 1.2 → NA として除外"],
        ["CLIP_LIMIT", "300", "runningFisher()", "各方向・合計スコアのクリッピング上限"],
        ["DEFAULT_P", "17,317", "runRFind()", "遺伝子ユニバースサイズ（暫定固定値）\n本来は commonGenes.length を使うべき"],
        ["バッチサイズ", "2", "runRFind()", "setTimeout バッチあたりの処理サンプル数"],
    ],
)

# 1.6 結果出力
doc.add_heading("1.6 結果出力", level=2)

doc.add_heading("1.6.1 結果テーブル", level=3)
doc.add_paragraph(
    "サンプル × インデックスのスコア行列をテーブル表示。"
    "Dx でグルーピング・ソート（Control 先頭）。"
    "Age 列が検出された場合は表示。"
    "スコアセルは色付き（正=緑系、負=赤系、scoreColor 関数で強度に応じた背景色）。"
)

doc.add_heading("1.6.2 Excel ダウンロード", level=3)
doc.add_paragraph("SheetJS を使用して .xlsx ファイルを生成。")
add_bullet("Summary シート: ", "SampleID, Dx, Demographics全列, 全インデックスのスコア")
add_bullet("インデックスごとのシート: ", "SampleID, Dx, Demographics全列, そのインデックスの RFind_Score")
doc.add_paragraph("シート名は最大31文字、禁止文字は _ に置換。Dx でソート済み。")

doc.add_heading("1.6.3 Figure 出力（Plotly.js）", level=3)
doc.add_paragraph("解析完了後、Figures セクションに以下のインタラクティブなグラフを自動生成。")

add_bullet("Box Plot（インデックスごと1枚）: ", "X軸=Dxグループ、Y軸=RFind Score\n"
           "Dx色分け、全データ点をjitter表示\nホバーで中央値・四分位表示")
add_bullet("Scatter Plot 個別（インデックスごと1枚、Age列がある場合のみ）: ",
           "X軸=Age、Y軸=RFind Score、Dx色分け")
add_bullet("Scatter Plot まとめ（全インデックス1枚、Age列＋インデックス2個以上の場合）: ",
           "2列グリッドのサブプロット配置\n凡例は1つ目のサブプロットのみ")
doc.add_paragraph("各 figure に PNG / SVG ダウンロードボタン付き。Plotly.downloadImage() による高解像度出力（scale=2）。")

doc.add_heading("1.6.4 Dx 色定義", level=3)
add_table(
    ["Dx", "色コード", "用途"],
    [
        ["Control", "#4ade80（緑）", "Box plot, Scatter, テーブル"],
        ["Schizo", "#f87171（赤）", "同上"],
        ["Bipolar", "#a78bfa（紫）", "同上"],
        ["MDD", "#fb923c（橙）", "同上"],
        ["その他", "#94a3b8（グレー）", "未定義の Dx 値"],
    ],
)

# 1.7 UI
doc.add_heading("1.7 UI / UX", level=2)

doc.add_heading("1.7.1 ステータスバー", level=3)
doc.add_paragraph(
    "ヘッダー下部に HGNC / Ortholog マッピングの読込状態を表示。"
    "3状態: loading（グレー）/ ✓ loaded（緑/青）/ ✗ not loaded（赤）。"
    "日本語の説明テキスト付き。"
    "HTTP サーバー経由でないと fetch が失敗し赤表示になる。"
)

doc.add_heading("1.7.2 中止ボタン", level=3)
doc.add_paragraph(
    "解析実行中に赤い「解析を中止」ボタンが表示される。"
    "クリックで _cancelFlag = true をセット、次のバッチで停止。"
    "UI をリセットし、結果は表示しない。"
)

doc.add_heading("1.7.3 i18n（国際化）", level=3)
doc.add_paragraph(
    "全 UI テキスト・エラーメッセージを日本語/英語で切替可能。"
    "T オブジェクトに ja/en の翻訳キーを格納、t(key) 関数で取得。"
    "alert() は使用せず、全て showError() による赤枠ボックス表示。"
)

# ════════════════════════════════════════════════════════
# 2. 作業内容
# ════════════════════════════════════════════════════════
doc.add_heading("2. 実装内容", level=1)

doc.add_heading("2.1 v2 → v3 変更対比", level=2)
add_table(
    ["項目", "v2（旧）", "v3（新）"],
    [
        ["入力", "DEG リスト（UP/DN 済み）", "FPKM マトリックス + Demographics"],
        ["参照インデックス", "プリセット6種（ハードコード）", "ユーザーがアップロード（複数可）"],
        ["FC 計算", "ユーザーが事前に計算", "Control 平均から自動計算"],
        ["Fisher 検定", "単純な 2×2 Fisher", "Running Fisher（累積 hypergeometric、双方向）"],
        ["結果", "1つの DEG vs 複数インデックス", "サンプル × インデックスのスコア行列"],
        ["ステップ数", "3ステップ", "4ステップ"],
        ["出力形式", "CSV + SVG (Forest Plot)", "Excel (.xlsx, 複数シート) + Plotly (PNG/SVG)"],
        ["可視化", "Forest Plot (SVG)", "Box plot + Scatter plot (Plotly.js)"],
    ],
)

doc.add_heading("2.2 実装した全機能一覧", level=2)

features = [
    ("Step 1: FPKM アップロード",
     "CSV/TSV/Excel 対応、遺伝子列ヘッダー検証、数値検証、HGNC+Ortholog正規化、プレビュー表示"),
    ("Step 2: Demographics アップロード",
     "ID/Dx/Age 列自動検出（列順不問）、FPKM との ID 照合、Control 検証、"
     "エラーメッセージ（日英）、不一致サンプルの除外リスト表示"),
    ("Step 3: インデックスアップロード",
     "複数ファイル一括対応、Symbol/FC 列自動検出（列順不問）、"
     "FC 形式自動判定（Ratio/Log2FC/SignedFC）・変換、HGNC+Ortholog正規化、個別削除"),
    ("Step 4: 解析実行",
     "Control 平均計算、FC 変換・フィルタリング、Running Fisher JS 移植、"
     "プログレスバー、中止ボタン、バッチ処理（UIブロック回避）"),
    ("Running Fisher 高速化",
     "hit-only スキャン最適化: phyperUpper をoverlap発生時のみ計算。"
     "263サンプル×1インデックスで 90秒→20秒"),
    ("結果テーブル",
     "サンプル × インデックスのスコア行列、色付きセル、Age 列表示、Dx ソート"),
    ("Excel 出力",
     "Summary シート + インデックスごとのシート、Demographics 全列含む"),
    ("Figure 出力（Plotly.js）",
     "Box plot（Dx別、インデックスごと）+ Scatter plot（Age vs Score、個別+まとめ）、"
     "PNG/SVG ダウンロード、インタラクティブ（ホバー、ズーム）"),
    ("HGNC 遺伝子名正規化",
     "97,737 エントリのマッピング、FPKM/Index 両方で適用、正規化数バッジ表示"),
    ("Ortholog マッピング",
     "Mouse/Rat → Human、2,042 エントリ（RGD データ）、FPKM/Index 両方対応"),
    ("FC 自動判定",
     "Ratio / Log2FC / SignedFC の3形式を自動検出、SignedFC に統一変換"),
    ("バリデーション",
     "各ステップで詳細なエラーメッセージ（日英対応、検出ヘッダー表示、ID例表示等）。"
     "alert() 不使用、全て showError() による赤枠ボックス"),
    ("i18n",
     "日本語/英語全テキスト切替、エラーメッセージ含む"),
    ("ステータスバー",
     "HGNC / Ortholog マッピングの読込状態を UI 表示（loading/loaded/error）"),
    ("中止ボタン",
     "解析中に赤い中止ボタン表示、クリックで停止、UI リセット"),
]
for title_text, desc in features:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(title_text + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("2.3 ファイル構成", level=2)
add_table(
    ["ファイル", "サイズ", "説明"],
    [
        ["index.html", "~55 KB", "メインアプリケーション（HTML + CSS + JS）"],
        ["hgnc_map.json", "2.1 MB", "HGNC alias→symbol マッピング（生成物）"],
        ["ortholog_map.json", "42 KB", "Mouse/Rat→Human ortholog マッピング（生成物）"],
        ["scripts/build_hgnc_map.py", "~2 KB", "HGNC マッピング生成（HGNC TSV → JSON）"],
        ["scripts/build_ortholog_map.py", "~2 KB", "Ortholog マッピング生成（RGD TSV → JSON）"],
        ["Chunk3_RFind.R", "~11 KB", "元の R スクリプト（参照用）"],
        ["Technotes_Illumina.pdf", "—", "Illumina Running Fisher アルゴリズム公式文書"],
    ],
)

doc.add_heading("2.4 主要関数一覧", level=2)
add_table(
    ["関数名", "目的"],
    [
        ["handleFPKMFile(file)", "FPKM ファイルの解析・検証・格納"],
        ["handleDemoFile(file)", "Demographics ファイルの解析・検証・格納"],
        ["matchDemoIfReady()", "FPKM と Demographics の ID 照合"],
        ["addIndexFromFile(file)", "インデックスファイルの解析・格納"],
        ["parseIndexCSV(aoa)", "インデックス CSV/Excel のパース（Symbol/FC列検出、FC変換）"],
        ["detectFCFormat(values)", "FC 形式自動判定（Ratio/Log2FC/SignedFC）"],
        ["convertFC(fc, format)", "FC を SignedFC に変換"],
        ["normalizeGene(sym)", "遺伝子名正規化（Ortholog → HGNC → uppercase）"],
        ["ensureHGNC()", "HGNC/Ortholog マッピングの非同期読込"],
        ["runRFind()", "メイン解析実行（FC計算→Running Fisher→結果格納）"],
        ["computeSampleFC(geneIdx, fpkm, controlMean)", "1遺伝子のFC計算（閾値・符号変換含む）"],
        ["runningFisher(b1, b2, P, clipLimit)", "Running Fisher コア実装"],
        ["phyperUpper(k, N, K, n)", "超幾何分布上側確率"],
        ["renderResults(res)", "結果テーブル＋Figure描画"],
        ["drawFigures(res)", "Plotly.js によるBox plot/Scatter plot生成"],
        ["dlExcel()", "Excel (.xlsx) ダウンロード（複数シート）"],
        ["cancelRFind()", "解析中止"],
        ["showError(msg) / clearError()", "エラー表示/消去"],
        ["updateStatusBar()", "ステータスバー更新"],
    ],
)

doc.add_heading("2.5 外部依存ライブラリ", level=2)
add_table(
    ["ライブラリ", "CDN URL", "用途", "サイズ目安"],
    [
        ["SheetJS", "cdn.sheetjs.com/xlsx-0.20.3/.../xlsx.full.min.js", "Excel/CSV 読み書き", "~300 KB"],
        ["Plotly.js", "cdn.jsdelivr.net/npm/plotly.js-dist-min@2", "インタラクティブ可視化", "~3 MB"],
        ["Google Fonts", "fonts.googleapis.com/css2?...", "Playfair Display, IBM Plex Sans/Mono", "~50 KB"],
    ],
)

# ════════════════════════════════════════════════════════
# 3. 残された課題
# ════════════════════════════════════════════════════════
doc.add_heading("3. 残された課題", level=1)

doc.add_heading("3.1 優先度: 高", level=2)
add_table(
    ["課題", "詳細", "備考"],
    [
        ["P 値（universe size）の適正化",
         "DEFAULT_P = 17,317 の固定値から commonGenes.length（FPKM∩Index 共通遺伝子数）に変更。"
         "Illumina公式定義に準拠。コード変更は1行だが、速度への影響を検証してから適用。",
         "一度実装したが速度問題で revert 済み"],
        ["数値検証",
         "R スクリプト（Chunk3_RFind.R）のローカル実行結果との数値照合。"
         "同一データで同一スコアが出ることを確認。",
         "通しテストは通過済み"],
    ],
)

doc.add_heading("3.2 優先度: 中", level=2)
add_table(
    ["課題", "詳細"],
    [
        ["10秒目標の高速化",
         "現状20秒（263サンプル×1インデックス）→10秒目標。Web Worker（マルチスレッド並列）の導入を検討"],
        ["結果テーブルのソート機能",
         "各列ヘッダーのクリックでソート切替"],
        ["ステータスバーの表示改善",
         "現状の表示がユーザーにとって意味がわかりにくい。UI/UX の改善"],
        ["Excel複数シート入力対応",
         "R スクリプトの indexes_all.xlsx のように、1つの Excel ファイル内の各シートを個別インデックスとして読み込み"],
    ],
)

doc.add_heading("3.3 優先度: 低", level=2)
add_table(
    ["課題", "詳細"],
    [
        ["デモデータ機能",
         "v2 にあったデモデータ機能の v3 対応（FPKM + Demographics + Index のサンプルセット）"],
        ["HGNC / Ortholog マッピングの定期更新",
         "HGNC/RGD データは定期更新される。build スクリプトの再実行フロー整備"],
        ["HTTP サーバー不要化",
         "file:// プロトコルでも JSON マッピングが読めるよう、"
         "マッピングを HTML にインライン埋め込みするオプション"],
    ],
)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), "..", "RFind_WebTool_260330.docx")
output_path = os.path.abspath(output_path)
doc.save(output_path)
print(f"Saved: {output_path}")
