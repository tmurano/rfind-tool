#!/usr/bin/env python3
"""Generate RFind Web Tool v3 design specification report as .docx — Final version 2026-03-30"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Style setup ──
style = doc.styles["Normal"]
font = style.font
font.name = "Yu Gothic"
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.4

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Yu Gothic"
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


# ════════════════════════════════════════════════════════
# Title
# ════════════════════════════════════════════════════════
title = doc.add_heading("RFind Web Tool v3 — 設計仕様書", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("最終更新: 2026-03-30", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("公開URL: https://tmurano.github.io/rfind-tool/", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# ════════════════════════════════════════════════════════
# 1. 設計と仕様書
# ════════════════════════════════════════════════════════
doc.add_heading("1. 設計と仕様書", level=1)

# 1.1 概要
doc.add_heading("1.1 ツール概要", level=2)
doc.add_paragraph(
    "RFind Web Tool v3 は、遺伝子発現量（FPKM/TPM）マトリックスと臨床情報（Demographics）を入力とし、"
    "参照インデックス（DEGリスト）との方向性一致度を Running Fisher 法（4方向）で定量化するクライアントサイド Web アプリケーションである。"
    "通常のpathway解析と異なり、遺伝子の発現変動の「方向（UP/DN）」まで考慮してスコアを算出する。"
    "すべての計算はブラウザ上で完結し、サーバーへのデータ送信は行わない。複数ユーザーが同時に使用しても互いに影響しない。"
)
doc.add_paragraph(
    "例えば「疾患群で発生関連遺伝子が胎児的パターンに回帰しているか？」「加齢で成熟マーカーが減少しているか？」"
    "といった研究質問に対して、サンプルごとのスコア・統計検定・論文用Figure・Methods文を一括で提供する。"
)

doc.add_heading("1.1.1 技術スタック", level=3)
add_table(
    ["項目", "内容"],
    [
        ["構成", "単一 HTML ファイル（CSS + JavaScript インライン、約95KB）"],
        ["外部ライブラリ", "SheetJS (xlsx-0.20.3): Excel 読み書き\n"
         "Plotly.js (v2): インタラクティブ可視化（Box plot, Scatter plot）"],
        ["対応言語", "日本語 / English（i18n 全テキスト切替）"],
        ["UI テーマ", "ダークテーマ（CSS カスタムプロパティ）"],
        ["計算方式", "全てクライアントサイド（ブラウザ内 JavaScript）"],
        ["公開方式", "GitHub Pages（静的ファイルホスティング）"],
        ["対応ブラウザ", "Chrome, Firefox, Safari, Edge（モダンブラウザ）"],
    ],
)

# 1.2 処理フロー
doc.add_heading("1.2 処理フロー（4ステップ + 結果出力）", level=2)

doc.add_heading("Step 1: 発現量マトリックスのアップロード", level=3)
doc.add_paragraph("遺伝子 × サンプルの発現量マトリックスをアップロードする。")
add_table(
    ["検証項目", "内容", "エラー時の挙動"],
    [
        ["ファイル形式", "CSV, TSV, Excel (.xlsx/.xls)", "対応外形式はパースエラー"],
        ["最小構造", "ヘッダー行 + データ行1行以上、2列以上", "エラー表示（日英対応）"],
        ["1列目ヘッダー", "Gene列として認識可能な名前\n（Gene, Symbol, GeneID, GeneName,\nEnsembl, Entrez, Probe, Name, ID 等）", "認識不可→エラー＋検出ヘッダー値を表示"],
        ["空サンプル名", "ヘッダー行に空のサンプル名がないこと", "エラー表示"],
        ["数値検証", "発現量の50%以上が非数値なら拒否", "エラー＋非数値セル数を表示"],
        ["遺伝子名正規化", "HGNC + Ortholog マッピング自動適用\n正規化数をプレビューバッジに表示", "マッピング未読込時はupperCase のみ"],
        ["発現量単位", "FPKM, TPM を想定（単位の検証は行わない）", "—"],
    ],
)

doc.add_heading("Step 2: サンプル情報（Demographics）のアップロード", level=3)
doc.add_paragraph('サンプル ID と診断群（Dx）を含むファイルをアップロード。Dx = "Control" が FC 算出の基準。')
add_table(
    ["検証項目", "内容", "エラー時の挙動"],
    [
        ["ID列", "ヘッダー名で自動検出（列順不問）\nID, SampleID, Sample_ID, SampleName,\nSubject, SubjectID, BrNum, Sample 等", "検出不可→エラー＋ヘッダー一覧表示"],
        ["Dx列", "ヘッダー名で自動検出（列順不問）\nDx, Diagnosis, Diag, Group, Condition 等", "検出不可→エラー＋ヘッダー一覧表示"],
        ["Age列", "自動検出（任意）\nAge, AgeDeath, Age_at_Death 等\n検出→結果テーブルとScatter plotに使用", "未検出でも正常動作"],
        ["Control 必須", 'Dx列に "Control" が最低1件必要', "0件→エラー＋検出Dx値一覧\n＋近い値（CON, ctrl等）があれば修正提案"],
        ["ID照合", "FPKM のサンプル ID と突合\n不一致サンプルは自動除外＋除外リスト表示", "0件マッチ→エラー＋両側ID例\n＋部分一致ヒントを表示"],
        ["マッチControl", "マッチしたサンプル中にControlが最低1件", "0件→エラー"],
        ["追加列", "ID・Dx以外の全列を保持\nExcel出力時に含まれる", "—"],
    ],
)

doc.add_heading("Step 3: 参照インデックス（Gene Set）", level=3)
doc.add_paragraph(
    "Running Fisher の参照シグネチャ（b1）となるDEGリストをアップロード、またはプリセットから選択。"
    "通常のpathway解析と異なり、遺伝子の発現変動の方向（UP/DN）まで考慮してスコアを算出する。"
)

doc.add_heading("プリセットインデックス", level=3)
doc.add_paragraph("以下の4つのプリセットインデックスがワンクリックで追加可能（サーバー上のCSVを自動読込）:")
add_table(
    ["名前", "説明", "ファイル", "遺伝子数"],
    [
        ["Dev", "胎生期→成熟の発生軸（Jaffe 2020）", "Dev.csv", "6,537"],
        ["imGC", "未成熟顆粒細胞マーカー（Zhou 2022）", "imGC_DEG.csv", "541"],
        ["Kainate D3", "カイニン酸誘発後3日目DEG（マウスDG）", "Kainate_Day03.csv", "2,654"],
        ["pH", "組織pH感受性遺伝子セット", "pH.csv", "304"],
    ],
)

doc.add_heading("カスタムアップロード", level=3)
add_table(
    ["検証項目", "内容", "エラー時の挙動"],
    [
        ["必須列", "Symbol列 + Fold Change列（列順不問）", "検出不可→エラー＋ヘッダー一覧＋許容列名例"],
        ["Symbol列", "Gene, Symbol, GeneName, GeneID,\nProbe, Name 等を許容", "—"],
        ["FC列", "Fold Change, FC, LogFC, Log2FC,\nFoldChange, Fold_Change 等を許容", "—"],
        ["P-Value/Rank列", "あっても無視（使用しない）", "—"],
        ["メタデータ行", "ヘッダー前のメタデータ行を自動スキップ\n（先頭20行からSymbol+FC列を含む行を検出）", "—"],
        ["FC形式自動判定", "Ratio / Log2FC / SignedFC を自動検出", "判定結果をpillに表示"],
        ["遺伝子名正規化", "HGNC + Ortholog マッピング適用", "—"],
        ["複数ファイル", "一括アップロード・D&D対応\n個別削除ボタンあり", "—"],
        ["Excel複数シート", "1つのExcelファイル内の各シートを\n個別インデックスとして自動読込\n（platform_info等の非データシートはスキップ）", "有効シート0件→エラー"],
    ],
)

doc.add_heading("1.2.1 FC 形式の自動判定ロジック", level=3)
add_table(
    ["形式", "値の範囲", "検出条件", "変換処理"],
    [
        ["Ratio（パターン1）", "0 ~ +∞", "全値が正", "FC<1 → -1/FC（DN）\nFC≥1 → そのまま（UP）"],
        ["Log2FC（パターン2）", "-∞ ~ +∞", "負値あり、|値|<1 が存在", "正 → 2^fc（UP）\n負 → -(2^|fc|)（DN）"],
        ["SignedFC（パターン3）", "-∞~-1, +1~+∞", "負値あり、(-1,+1)内の値なし", "変換なし"],
    ],
)
doc.add_paragraph(
    "注意: 全値正のDEGリスト（UP方向のみ）はRatio判定される。FC≥1なら全てUP方向として扱われ、"
    "Running FisherのUU/UD方向でのみスコアが出る（DN方向は0）。これは正常な動作。"
)

doc.add_heading("Step 4: 解析実行", level=3)
doc.add_paragraph("サンプル × インデックスごとに以下のパイプラインを実行。")

doc.add_paragraph("(a) Control 平均 FPKM の算出", style="List Number")
doc.add_paragraph("各遺伝子について、Dx = Control の全サンプルの FPKM 平均を算出。")

doc.add_paragraph("(b) Fold Change (FC) の計算", style="List Number")
doc.add_paragraph("FC = サンプル FPKM ÷ Control 平均 FPKM")

doc.add_paragraph("(c) FC のフィルタリングと符号変換", style="List Number")
add_table(
    ["条件", "処理"],
    [
        ["Control平均 = 0 または NaN", "NA（除外）"],
        ["|FC| < 1.2（閾値: THRESHOLD_FC = 1.2）", "NA（除外）"],
        ["FC ≤ 1", "-1/FC に変換（負の符号 = downregulated）"],
        ["FC > 1", "そのまま（正 = upregulated）"],
    ],
)

doc.add_paragraph("(d) Running Fisher スコアの算出", style="List Number")
doc.add_paragraph("b1（参照インデックス）と b2（サンプルの FC）で Running Fisher を実行。b2 の有効遺伝子数が 5 未満の場合はスコア = 0。")

doc.add_paragraph("(e) バッチ処理と中止機能", style="List Number")
doc.add_paragraph(
    "2サンプルずつ setTimeout でバッチ処理し、UI をブロックしない。"
    "プログレスバーで進捗表示。赤い「解析を中止」ボタンでキャンセル可能。"
)

# 1.3 Running Fisher
doc.add_heading("1.3 Running Fisher アルゴリズム詳細", level=2)

doc.add_paragraph(
    "Illumina BaseSpace Correlation Engine の Running Fisher アルゴリズム（Tech Note 970-2014-007）に準拠。"
    "GSEA と類似のランクベース enrichment だが、統計的有意性を Fisher's exact test で評価する点が異なる。"
)

doc.add_heading("1.3.1 入力", level=3)
add_table(
    ["パラメータ", "説明"],
    [
        ["b1", "参照インデックスの遺伝子リスト [{Symbol, FC}]"],
        ["b2", "サンプルの FC 変換済み遺伝子リスト [{Symbol, FC}]"],
        ["P", "遺伝子ユニバースサイズ = FPKM ∩ Index の共通遺伝子数\n（Illumina定義: P1∩P2）\nインデックスごとに自動計算"],
        ["CLIP_LIMIT", "スコアクリッピング上限 = 300"],
    ],
)

doc.add_heading("1.3.2 アルゴリズムの流れ", level=3)

steps = [
    "b1, b2 それぞれを重複除去（Symbol の lowercase で一意化）し、FC=0 を除外",
    "各セットを UP (FC>0) と DN (FC<0) に分割",
    "4方向の組み合わせで fisherBidirectional を実行: UU, UD, DU, DD",
    "fisherBidirectional の内部処理（方向 dir1, dir2 について）:\n"
    "  a. b1 から dir1 方向の遺伝子を FC 順にランク付け（ranked1）\n"
    "  b. b2 から dir2 方向の遺伝子を FC 順にランク付け（ranked2）\n"
    "  c. Direction 1: ranked1 を上位からスキャン、ranked2 との累積重複を数える\n"
    "     → hit があった時のみ phyperUpper(cum, P, K, n) を計算（高速化最適化）\n"
    "     → 全ランク中の最小 p 値 (p1) と最適ランク (m1) を記録\n"
    "  d. Direction 2: ranked2 を逆方向にスキャンして同様に p2, m2 を取得\n"
    "  e. 双方向の平均: meanLogP = (-log10(p1) + -log10(p2)) / 2",
    "各方向のスコア = -log10(p) × sign（concordant: +1, discordant: -1）\n"
    "  各方向をクリッピング（±CLIP_LIMIT）",
    "Total = 4方向スコアの合計 → さらにクリッピング（±CLIP_LIMIT）",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}", style="List Number")

doc.add_heading("1.3.3 方向と符号", level=3)
add_table(
    ["方向", "b1", "b2", "符号", "意味"],
    [
        ["UU", "UP", "UP", "+1", "Concordant（方向一致）"],
        ["UD", "UP", "DOWN", "-1", "Discordant"],
        ["DU", "DOWN", "UP", "-1", "Discordant"],
        ["DD", "DOWN", "DOWN", "+1", "Concordant（方向一致）"],
    ],
)

doc.add_heading("1.3.4 P値（ユニバースサイズ）の定義", level=3)
doc.add_paragraph(
    "Illumina Tech Note (Figure 4) に準拠し、P = P1∩P2（FPKMプラットフォームとインデックスプラットフォームの共通遺伝子数）を使用。"
    "インデックスごとに異なるPが自動計算される。"
    "commonGenes = FPKM.genes ∩ Index.genes がPに相当する。"
)

doc.add_heading("1.3.5 phyperUpper の高速化（hit-only スキャン）", level=3)
doc.add_paragraph(
    "Running Fisher のスキャンでは、各ランクで phyperUpper（超幾何分布の上側確率）を計算する。"
    "以下の数学的性質を利用して高速化："
)
add_bullet("hit なし（overlap 不変）の場合: n が増加、overlap 同一 → p 値は必ず悪化 → スキップ可能", "")
add_bullet("hit あり（overlap 増加）の場合のみ: p 値が改善する可能性 → 計算実行", "")
doc.add_paragraph(
    "この最適化により phyperUpper の呼び出し回数が「全遺伝子数（数千回）」から"
    "「実際の overlap 数（数十回）」に激減。263サンプル×4インデックスで約18秒。"
    "結果は数学的に同一。"
)

# 1.4 HGNC + Ortholog
doc.add_heading("1.4 遺伝子名正規化", level=2)

doc.add_heading("1.4.1 HGNC Alias → Official Symbol", level=3)
add_table(
    ["項目", "内容"],
    [
        ["ファイル", "hgnc_map.json（2.1 MB）"],
        ["エントリ数", "97,737"],
        ["対象", "Previous symbols, Alias symbols, Ensembl Gene ID → Approved Symbol"],
        ["Ambiguous alias", "1,654件を除外（1つの alias が複数の approved symbol に対応するもの）"],
        ["生成スクリプト", "scripts/build_hgnc_map.py\nHGNC カスタムダウンロードから TSV 取得→JSON 変換"],
        ["読込方式", "ページ読込時に非同期 fetch（HTTPサーバー必須）"],
    ],
)

doc.add_heading("1.4.2 Ortholog マッピング（Mouse/Rat → Human）", level=3)
add_table(
    ["項目", "内容"],
    [
        ["ファイル", "ortholog_map.json（42 KB）"],
        ["エントリ数", "2,042（uppercase で異なる ortholog のみ）"],
        ["データソース", "RGD_ORTHOLOGS_Ensembl.txt（Mouse + Rat → Human を1ファイルで取得）"],
        ["Ambiguous", "11件を除外"],
        ["生成スクリプト", "scripts/build_ortholog_map.py"],
    ],
)

doc.add_heading("1.4.3 normalizeGene() の処理順序", level=3)
doc.add_paragraph("1. 入力を uppercase 化", style="List Number")
doc.add_paragraph("2. Ortholog マッピングで検索（Mouse/Rat → Human）→ ヒットすれば変換", style="List Number")
doc.add_paragraph("3. HGNC マッピングで検索（alias → official symbol）→ ヒットすれば変換", style="List Number")
doc.add_paragraph("4. どちらにもヒットしなければ uppercase のまま返す", style="List Number")

# 1.5 パラメータ
doc.add_heading("1.5 パラメータ一覧", level=2)
add_table(
    ["パラメータ", "値", "備考"],
    [
        ["THRESHOLD_FC", "1.2", "FC 閾値（|FC| < 1.2 → NA として除外）"],
        ["CLIP_LIMIT", "300", "各方向・合計スコアのクリッピング上限"],
        ["P (Universe)", "commonGenes.length", "インデックスごとにFPKM∩Indexの共通遺伝子数を自動計算（Illumina仕様）"],
        ["バッチサイズ", "2", "setTimeout バッチあたりの処理サンプル数"],
    ],
)

# 1.6 結果出力
doc.add_heading("1.6 結果出力", level=2)

doc.add_heading("1.6.1 結果テーブル", level=3)
doc.add_paragraph(
    "サンプル × インデックスのスコア行列をテーブル表示。"
    "Dx でグルーピング・ソート（Control 先頭）。Age 列は自動検出時に表示。"
    "スコアセルは色付き（正=緑系、負=赤系）。"
    "列ヘッダークリックで昇順/降順ソート。"
    "スコアセルクリックで Overlapping genes 詳細パネルを表示。"
    "各インデックス列ヘッダーに「?」ツールチップ（スコアの解釈方法を説明）。"
)

doc.add_heading("1.6.2 Overlapping genes 詳細", level=3)
doc.add_paragraph(
    "結果テーブルのスコアセルをクリックすると、そのサンプル×インデックスの詳細パネルが開く。"
    "4方向（UU, DD, UD, DU）ごとにスコア、overlap数、p値、寄与遺伝子リストを表示。"
    "各遺伝子タグにホバーすると Index FC と Sample FC が表示される。"
    "クリック時に1回だけ再計算（returnGenes=true）するのでメモリ消費なし。"
)

doc.add_heading("1.6.3 Excel ダウンロード", level=3)
doc.add_paragraph("SheetJS を使用して .xlsx ファイルを生成。")
add_bullet("Summary シート: ", "SampleID, Dx, Demographics全列, 全インデックスのスコア")
add_bullet("インデックスごとのシート: ", "SampleID, Dx, Demographics全列, そのインデックスの RFind_Score")
add_bullet("Parameters シート: ", "日時, ファイル名, 遺伝子数, サンプル数, Control数,\n"
           "インデックスごとの遺伝子数・共通遺伝子数(P)・FC形式,\n"
           "THRESHOLD_FC, CLIP_LIMIT, アルゴリズム名, 参照文献")

doc.add_heading("1.6.4 Figure 出力（Plotly.js）", level=3)
add_bullet("Box Plot（インデックスごと1枚）: ", "X軸=Dxグループ、Y軸=RFind Score\n"
           "Dx色分け、全データ点jitter表示\np値（Wilcoxon rank-sum test vs Control）をアノテーション表示\n"
           "「?」ツールチップで解釈方法を説明")
add_bullet("Scatter Plot 個別（インデックスごと1枚、Age列がある場合のみ）: ",
           "X軸=Age、Y軸=RFind Score、Dx色分け\n"
           "全サンプルの線形回帰直線（点線）＋R²＋p値を凡例に表示\n"
           "「?」ツールチップで解釈方法を説明")
add_bullet("Scatter Plot まとめ（全インデックス1枚、Age＋2インデックス以上）: ",
           "2列グリッドのサブプロット配置\n各サブプロットにインデックス名＋R²＋p値\n回帰直線付き")

doc.add_heading("1.6.5 Figure Settings", level=3)
doc.add_paragraph("Figuresセクションの「Figure Settings」ボタンから設定パネルを開く。")
add_table(
    ["設定", "デフォルト", "説明"],
    [
        ["Width / Height", "1000px", "ダウンロード画像サイズ（正方形）"],
        ["Scale (DPI)", "2x (144dpi)", "1x〜4xから選択"],
        ["Font size", "14", "軸ラベル等のフォントサイズ"],
        ["Title font", "16", "タイトルのフォントサイズ"],
        ["Font family", "Arial", "Arial / Helvetica / Times New Roman / IBM Plex Sans"],
        ["Axis label X/Y", "auto", "空欄→デフォルト、入力→上書き"],
        ["Marker size", "7", "Scatter のドットサイズ"],
        ["Dx colors", "各色", "カラーピッカーで各Dxの色を変更"],
    ],
)
doc.add_paragraph("「Apply & Redraw」で全Figureに反映。PNG/SVGダウンロードも設定値で出力。")

doc.add_heading("1.6.6 Methods文自動生成", level=3)
doc.add_paragraph(
    "「Methods文を生成」ボタンで、論文のMethodsセクションにそのまま使える英文を生成。"
    "解析パラメータ（サンプル数、Control数、インデックス名、Dx群内訳、P値の定義等）が自動埋め込み。"
    "引用文献（Kupershmidt et al., 2010; Illumina Tech Note）も含む。"
    "自動的にクリップボードにコピーされる。"
)

doc.add_heading("1.6.7 統計検定", level=3)
add_table(
    ["検定", "用途", "実装"],
    [
        ["Wilcoxon rank-sum test\n（Mann-Whitney U）", "Box plot: 各Dx vs Controlのp値", "正規近似（normalCDF）\n***p<0.001 **p<0.01 *p<0.05 ns"],
        ["線形回帰 + t検定", "Scatter plot: Age vs RFscoreの\n回帰直線、R²、p値", "最小二乗法\nslope の t検定"],
    ],
)

doc.add_heading("1.6.8 Dx 色定義", level=3)
add_table(
    ["Dx", "色コード"],
    [
        ["Control", "#4ade80（緑）"],
        ["Schizo", "#f87171（赤）"],
        ["Bipolar", "#a78bfa（紫）"],
        ["MDD", "#fb923c（橙）"],
        ["その他", "#94a3b8（グレー）"],
    ],
)

# 1.7 UI/UX
doc.add_heading("1.7 UI / UX", level=2)

doc.add_heading("1.7.1 デモデータ", level=3)
doc.add_paragraph(
    "「デモデータで試す」ボタンでサンプルデータ（FPKM_demo.xlsx + Demographics_demo.xlsx + Dev.csv + imGC_DEG.csv）を"
    "自動読込→解析実行→結果表示まで一気に実行。ユーザーは結果のイメージを即座に確認可能。"
    "デモデータの差し替えはファイルを置き換えるだけでHTMLの編集不要。"
)

doc.add_heading("1.7.2 ヘルプツールチップ", level=3)
doc.add_paragraph(
    "結果テーブルの各インデックス列ヘッダー、Box plot タイトル、Scatter plot タイトルの横に「?」アイコン。"
    "ホバーでスコアの解釈方法を日本語/英語で表示。例："
)
add_bullet("スコア列: 「正のスコア→同じ方向に変動、負→逆方向。例えばDevで正なら胎児的パターン」", "")
add_bullet("Box plot: 「Control群と比較してスコアが高い群はそのインデックスのパターンに近い」", "")
add_bullet("Scatter: 「傾きが正→年齢とともにパターン強化、負→弱化」", "")

doc.add_heading("1.7.3 入力テンプレート", level=3)
doc.add_paragraph(
    "各Stepのガイドボックスに「テンプレート」ボタン。"
    "クリックでExcelテンプレート（サンプルデータ付き）がダウンロードされる。"
)

doc.add_heading("1.7.4 エラーサジェスト", level=3)
doc.add_paragraph(
    "Control未検出時: 近い値（CON, ctrl等）を検出して修正提案を表示。"
    "ID不一致時: 部分一致チェックで原因ヒント（接頭辞の違い or 完全に異なる命名規則）を表示。"
    "alert()は一切使用せず、全てshowError()による赤枠ボックス表示（日英対応）。"
)

doc.add_heading("1.7.5 中止ボタン", level=3)
doc.add_paragraph("解析実行中に赤い「解析を中止」ボタン。_cancelFlag で次バッチで停止。")

doc.add_heading("1.7.6 i18n（国際化）", level=3)
doc.add_paragraph("全UIテキスト・エラーメッセージ・ツールチップを日本語/英語で切替可能。T オブジェクトに翻訳キーを格納。")

# ════════════════════════════════════════════════════════
# 2. 実装内容
# ════════════════════════════════════════════════════════
doc.add_heading("2. 実装内容", level=1)

doc.add_heading("2.1 実装した全機能一覧", level=2)

features = [
    ("Step 1: 発現量マトリックスアップロード",
     "CSV/TSV/Excel対応、遺伝子列ヘッダー検証、数値検証、HGNC+Ortholog正規化、プレビュー表示"),
    ("Step 2: サンプル情報アップロード",
     "ID/Dx/Age列自動検出（列順不問）、FPKM との ID照合、Control検証、"
     "エラーサジェスト（日英）、不一致サンプルの除外リスト表示"),
    ("Step 3: 参照インデックス",
     "プリセット4種（Dev, imGC, Kainate D3, pH）ワンクリック追加\n"
     "カスタムアップロード（複数ファイル一括、D&D対応）\n"
     "Symbol/FC列自動検出（列順不問）、メタデータ行自動スキップ\n"
     "FC形式自動判定（Ratio/Log2FC/SignedFC）・変換\n"
     "Excel複数シート対応、HGNC+Ortholog正規化、個別削除"),
    ("Step 4: 解析実行",
     "Control平均計算、FC変換・フィルタリング、Running Fisher JS移植\n"
     "P=commonGenes.length（Illumina仕様）\n"
     "プログレスバー、中止ボタン、バッチ処理（UIブロック回避）"),
    ("Running Fisher 高速化",
     "hit-only スキャン: phyperUpper をoverlap発生時のみ計算\n"
     "263サンプル×4インデックスで約18秒"),
    ("結果テーブル",
     "サンプル×インデックスのスコア行列、色付きセル、Age列表示\n"
     "列ヘッダークリックソート、「?」ツールチップ"),
    ("Overlapping genes 詳細",
     "スコアセルクリックで4方向別の寄与遺伝子リスト表示\n"
     "各遺伝子のIndex FC / Sample FCをホバー表示"),
    ("Excel出力",
     "Summaryシート + インデックスごとシート + Parametersシート\n"
     "Demographics全列含む"),
    ("Figure出力（Plotly.js）",
     "Box plot（Dx別、p値付き）\n"
     "Scatter plot（Age vs Score、回帰直線+R²+p値、個別+まとめ）\n"
     "PNG/SVGダウンロード（正方形1000x1000、scale可変）\n"
     "Figure Settings（サイズ、DPI、フォント、色、マーカー）"),
    ("Methods文自動生成",
     "論文Methodsにそのまま使える英文を生成＋クリップボードコピー\n"
     "解析パラメータ・引用文献を自動埋め込み"),
    ("ヘルプツールチップ",
     "スコア列、Box plot、Scatter plotに「?」アイコン\n"
     "ホバーで解釈方法を日英で説明"),
    ("デモデータ",
     "ワンクリックでデモデータ読込→解析実行→結果表示\n"
     "FPKM_demo.xlsx + Demographics_demo.xlsx + Dev.csv + imGC_DEG.csv\n"
     "データ差し替えはファイル置換のみ"),
    ("HGNC遺伝子名正規化",
     "97,737エントリ（alias→official symbol + ENSEMBL ID）\nFPKM/Index両方で適用"),
    ("Orthologマッピング",
     "Mouse/Rat→Human、2,042エントリ（RGDデータ）\nFPKM/Index両方対応"),
    ("入力テンプレート",
     "各StepにExcelテンプレートダウンロードボタン"),
    ("エラーサジェスト",
     "Control未検出時の修正提案、ID不一致時の原因ヒント"),
    ("プリセットインデックス",
     "Dev, imGC, Kainate D3, pH をワンクリック追加\n"
     "CSVファイル追加＋PRESETS配列1行追加で拡張可能"),
    ("バリデーション",
     "各ステップで詳細エラー（日英対応）、alert()不使用"),
    ("i18n", "日本語/英語全テキスト切替"),
    ("中止ボタン", "解析中に中止可能、UIリセット"),
]
for title_text, desc in features:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(title_text + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("2.2 ファイル構成", level=2)
add_table(
    ["ファイル", "サイズ", "説明"],
    [
        ["index.html", "~95 KB", "メインアプリケーション（HTML + CSS + JS）"],
        ["hgnc_map.json", "2.1 MB", "HGNC alias→symbol マッピング（生成物）"],
        ["ortholog_map.json", "42 KB", "Mouse/Rat→Human ortholog マッピング（生成物）"],
        ["Dev.csv", "198 KB", "プリセットインデックス: 発生軸（6,537遺伝子）"],
        ["imGC_DEG.csv", "15 KB", "プリセットインデックス: imGCマーカー（541遺伝子）"],
        ["Kainate_Day03.csv", "100 KB", "プリセットインデックス: カイニン酸DEG（2,654遺伝子）"],
        ["pH.csv", "15 KB", "プリセットインデックス: pH感受性（304遺伝子）"],
        ["FPKM_demo.xlsx", "8.9 MB", "デモデータ: 発現量マトリックス（50サンプル）"],
        ["Demographics_demo.xlsx", "16 KB", "デモデータ: サンプル情報"],
        ["scripts/build_hgnc_map.py", "3 KB", "HGNC マッピング生成スクリプト"],
        ["scripts/build_ortholog_map.py", "3 KB", "Ortholog マッピング生成スクリプト"],
    ],
)

doc.add_heading("2.3 主要関数一覧", level=2)
add_table(
    ["関数名", "目的"],
    [
        ["handleFPKMFile(file)", "FPKMファイルの解析・検証・格納"],
        ["handleDemoFile(file)", "Demographicsファイルの解析・検証・格納"],
        ["matchDemoIfReady()", "FPKMとDemographicsのID照合"],
        ["addIndexFromFile(file)", "インデックスファイルの解析・格納（Excel複数シート対応）"],
        ["parseIndexCSV(aoa)", "インデックスCSV/Excelのパース（ヘッダー自動検出、FC変換）"],
        ["detectFCFormat(values)", "FC形式自動判定（Ratio/Log2FC/SignedFC）"],
        ["convertFC(fc, format)", "FCをSignedFCに変換"],
        ["normalizeGene(sym)", "遺伝子名正規化（Ortholog→HGNC→uppercase）"],
        ["ensureHGNC()", "HGNC/Orthologマッピングの非同期読込"],
        ["runRFind()", "メイン解析実行（FC計算→Running Fisher→結果格納）"],
        ["computeSampleFC(geneIdx, fpkm, controlMean)", "1遺伝子のFC計算（閾値・符号変換含む）"],
        ["runningFisher(b1, b2, P, clipLimit, returnGenes)", "Running Fisherコア実装（遺伝子リスト返却オプション付き）"],
        ["phyperUpper(k, N, K, n)", "超幾何分布上側確率"],
        ["mannWhitneyU(a, b)", "Wilcoxon rank-sum test（正規近似）"],
        ["linearRegression(xs, ys)", "線形回帰（slope, intercept, R², p）"],
        ["renderResults(res)", "結果テーブル描画（ソート機能付き）"],
        ["drawFigures(res)", "Plotly.jsによるBox plot/Scatter plot生成"],
        ["showOverlap(sampleId, idxIdx)", "Overlapping genes詳細パネル表示"],
        ["dlExcel()", "Excelダウンロード（複数シート+Parameters）"],
        ["copyMethods()", "Methods文生成＋クリップボードコピー"],
        ["loadDemoAndRun()", "デモデータ読込→自動解析実行"],
        ["loadPreset(i)", "プリセットインデックス読込"],
        ["cancelRFind()", "解析中止"],
        ["showError(msg) / clearError()", "エラー表示/消去"],
        ["dlTemplate(type)", "入力テンプレートExcelダウンロード"],
        ["applyFigSettings()", "Figure設定適用＋再描画"],
    ],
)

doc.add_heading("2.4 外部依存ライブラリ", level=2)
add_table(
    ["ライブラリ", "CDN URL", "用途"],
    [
        ["SheetJS", "cdn.sheetjs.com/xlsx-0.20.3/.../xlsx.full.min.js", "Excel/CSV 読み書き"],
        ["Plotly.js", "cdn.jsdelivr.net/npm/plotly.js-dist-min@2", "インタラクティブ可視化"],
        ["Google Fonts", "fonts.googleapis.com/css2?...", "Playfair Display, IBM Plex Sans/Mono"],
    ],
)

# ════════════════════════════════════════════════════════
# 3. 残された課題
# ════════════════════════════════════════════════════════
doc.add_heading("3. 残された課題", level=1)

doc.add_heading("3.1 優先度: 高", level=2)
add_table(
    ["課題", "詳細"],
    [
        ["数値検証",
         "Rスクリプト（Chunk3_RFind.R）のローカル実行結果との数値照合。"
         "同一データで同一スコアが出ることを確認。通しテストは通過済み。"],
    ],
)

doc.add_heading("3.2 優先度: 中〜低", level=2)
add_table(
    ["課題", "詳細"],
    [
        ["プリセットインデックスの拡充",
         "新しいDEGリストの追加（CSVファイル追加＋PRESETS配列1行追加で対応可能）"],
        ["HGNC / Ortholog マッピングの定期更新",
         "build スクリプトの再実行フロー整備"],
        ["ライトテーマ対応",
         "CSS変数の切替で対応可能だが未実装"],
    ],
)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), "..", "RFind_Web_260330.docx")
output_path = os.path.abspath(output_path)
doc.save(output_path)
print(f"Saved: {output_path}")
