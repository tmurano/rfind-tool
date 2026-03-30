#!/usr/bin/env python3
"""Generate RFind Web Tool v3 design specification report — v3 2026-03-30"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
style = doc.styles["Normal"]
font = style.font; font.name = "Yu Gothic"; font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.4
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]; hs.font.name = "Yu Gothic"; hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]; cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)

def bul(text, bp=None):
    p = doc.add_paragraph(style="List Bullet")
    if bp: run = p.add_run(bp); run.bold = True; p.add_run(text)
    else: p.add_run(text)

# Title
title = doc.add_heading("RFind Web Tool v3 — 設計仕様書 v3", level=0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("最終更新: 2026-03-30").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("公開URL: https://tmurano.github.io/rfind-tool/").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("GitHub: https://github.com/tmurano/rfind-tool").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# ═══ 1. 設計と仕様書 ═══
doc.add_heading("1. 設計と仕様書", level=1)

doc.add_heading("1.1 ツール概要", level=2)
doc.add_paragraph(
    "RFind Web Tool v3 は、遺伝子発現量（FPKM/TPM）マトリックスとサンプル情報を入力とし、"
    "参照インデックス（DEGリスト）との方向性一致度を Running Fisher 法（4方向）で定量化するクライアントサイド Web アプリケーション。"
    "通常のpathway解析と異なり、遺伝子の発現変動の「方向（UP/DN）」まで考慮してスコアを算出する。")
doc.add_paragraph(
    "キャッチコピー: あなたの遺伝子発現データと参照gene setの「方向を含めた類似度」を、サンプルごとに数値化します。"
    "すべての計算はブラウザ上で完結し、データは一切外部サーバーに送信されない。")

doc.add_heading("1.1.1 技術スタック", level=3)
add_table(["項目", "内容"], [
    ["構成", "単一HTMLファイル（CSS+JSインライン、約100KB）"],
    ["外部ライブラリ", "SheetJS (xlsx-0.20.3): Excel読み書き\nPlotly.js (v2): Box plot, Scatter plot"],
    ["対応言語", "日本語 / English（i18n全テキスト切替）"],
    ["UI", "ダークテーマ、accordion折りたたみ"],
    ["計算", "全てクライアントサイド（ブラウザ内JavaScript）"],
    ["公開", "GitHub Pages"],
])

# 1.2 処理フロー
doc.add_heading("1.2 処理フロー", level=2)

doc.add_heading("Step 1: 発現量マトリックス", level=3)
add_table(["検証項目", "内容", "エラー時"], [
    ["ファイル形式", "CSV, TSV, Excel", "パースエラー"],
    ["最小構造", "ヘッダー+データ1行、2列以上", "エラー(日英)"],
    ["1列目ヘッダー", "Gene列として認識可能な名前\n(Gene,Symbol,GeneID,Ensembl等)", "エラー+検出値"],
    ["空サンプル名", "なし", "エラー"],
    ["数値検証", "50%以上非数値→拒否", "エラー+セル数"],
    ["遺伝子名正規化", "HGNC+Ortholog自動適用", "未読込→upperCaseのみ"],
])

doc.add_heading("Step 2: サンプル情報", level=3)
doc.add_paragraph('Dx="Control"必須。理由: Control群の平均発現量を基準として各サンプルのFold Changeを計算し、遺伝子の変動方向を定義するため。')
add_table(["検証項目", "内容", "エラー時"], [
    ["ID列", "自動検出(列順不問)\nID,SampleID,Sample,Subject,BrNum等", "エラー+ヘッダー一覧"],
    ["Dx列", "自動検出(列順不問)\nDx,Diagnosis,Group,Condition等", "エラー+ヘッダー一覧"],
    ["Age列等", "自動検出(任意)。Figure横軸に使用可能", "未検出→正常動作"],
    ["Control必須", '"Control"最低1件', "エラー+Dx値一覧+修正提案\n(CON,ctrl等検出時)"],
    ["ID照合", "FPKMと突合、不一致は自動除外", "0件→エラー+ID例+ヒント"],
    ["追加列", "全列保持→Excel出力・Figure横軸に含む", "—"],
])

doc.add_heading("Step 3: 参照インデックス（Gene Set）", level=3)
doc.add_paragraph("DEGリストをアップロード、またはプリセットから選択。Gene Group（FC列なし）にも対応。")

doc.add_heading("プリセットインデックス", level=3)
add_table(["名前", "説明", "出典", "遺伝子数"], [
    ["Dev", "胎生期→成熟の発生軸", "Kang et al. 2011, Nature", "6,537"],
    ["imGC", "未成熟顆粒細胞マーカー", "Zhou et al. 2022, Nature", "541"],
    ["Excitability", "神経興奮性インデックス", "Mori et al. 2017, Sci Data", "2,654"],
    ["pH", "組織pH感受性遺伝子セット", "Atz et al. 2010, Am J Med Genet", "304"],
])
doc.add_paragraph("各プリセットに[ref]リンク付き。CSVファイル追加+PRESETS配列1行で拡張可能。")

doc.add_heading("カスタムアップロード", level=3)
add_table(["項目", "内容"], [
    ["必須列", "Symbol列（列順不問）。FC列は任意"],
    ["FC列あり", "FC形式自動判定(Ratio/Log2FC/SignedFC)"],
    ["FC列なし(Gene Group)", "全FC=1(UP方向)。方向性なしenrichment"],
    ["メタデータ行", "ヘッダー前を自動スキップ(先頭20行スキャン)"],
    ["Excel複数シート", "各シートを個別インデックスとして読込"],
])

doc.add_heading("1.2.1 FC形式の自動判定", level=3)
add_table(["形式", "検出条件", "変換"], [
    ["Ratio", "全値が正", "FC<1→-1/FC, FC≥1→そのまま"],
    ["Log2FC", "負値あり、|値|<1あり", "正→2^fc, 負→-(2^|fc|)"],
    ["SignedFC", "負値あり、(-1,+1)内なし", "変換なし"],
    ["Gene Group", "FC列なし", "全FC=1"],
])

doc.add_heading("Step 4: 解析実行", level=3)
doc.add_paragraph("(a) Control平均FPKM → (b) FC=FPKM/Control平均 → (c) |FC|<1.2→NA, FC≤1→-1/FC → (d) Running Fisher → (e) バッチ処理+中止ボタン")

# 1.3 Running Fisher
doc.add_heading("1.3 Running Fisher アルゴリズム", level=2)
doc.add_paragraph("Illumina BaseSpace Correlation Engine (Tech Note 970-2014-007) に準拠。")

doc.add_heading("1.3.1 流れ", level=3)
for i, s in enumerate([
    "b1,b2を重複除去(lowercase一意化)、FC=0除外",
    "UP(FC>0)とDN(FC<0)に分割",
    "4方向(UU,UD,DU,DD)でfisherBidirectional実行",
    "各方向: ranked1スキャン→hitのみphyperUpper計算→最小p値→逆方向も→双方向平均",
    "方向スコア = -log10(p)×sign (concordant:+1, discordant:-1)、クリップ±300",
    "Total = 4方向合計→クリップ±300",
], 1):
    doc.add_paragraph(f"{i}. {s}", style="List Number")

add_table(["方向","b1","b2","符号","意味"], [
    ["UU","UP","UP","+1","Concordant"],["UD","UP","DOWN","-1","Discordant"],
    ["DU","DOWN","UP","-1","Discordant"],["DD","DOWN","DOWN","+1","Concordant"],
])

doc.add_heading("1.3.2 P値", level=3)
doc.add_paragraph("P = FPKM∩Indexの共通遺伝子数(Illumina P1∩P2)。インデックスごとに自動計算。")

doc.add_heading("1.3.3 hit-only最適化", level=3)
doc.add_paragraph("hitなし→skip、hitあり→計算。263sample×4index≈18秒。結果同一。")

# 1.4 遺伝子名正規化
doc.add_heading("1.4 遺伝子名正規化", level=2)
add_table(["マッピング","ファイル","エントリ","ソース"], [
    ["HGNC(alias→official)","hgnc_map.json(2.1MB)","97,737","HGNC TSV"],
    ["Ortholog(Mouse/Rat→Human)","ortholog_map.json(42KB)","2,042","RGD"],
])
doc.add_paragraph("順序: uppercase → Ortholog → HGNC → そのまま。FPKM/Index両方で適用。")

# 1.5 パラメータ
doc.add_heading("1.5 パラメータ", level=2)
add_table(["パラメータ","値","備考"], [
    ["THRESHOLD_FC","1.2","|FC|<1.2→NA除外"],
    ["CLIP_LIMIT","300","スコアクリッピング±300"],
    ["P","commonGenes.length","インデックスごと自動計算"],
    ["バッチ","2 samples/batch","setTimeout処理"],
])

# 1.6 結果出力
doc.add_heading("1.6 結果出力", level=2)

doc.add_heading("1.6.1 テーブル", level=3)
doc.add_paragraph(
    "スコア行列、Dxソート、Age列表示、色付きセル、列ソート。"
    "各インデックス列に「?」ツールチップ（スコアの具体的な解釈方法＋例を日英で表示）。"
    "セルクリック→Overlapping genes詳細(4方向別の寄与遺伝子リスト、Index FC/Sample FC表示)。")

doc.add_heading("1.6.2 Excel", level=3)
doc.add_paragraph("Summary + インデックスごと + Parametersシート。Demographics全列含む。")

doc.add_heading("1.6.3 Figure(Plotly.js)", level=3)
doc.add_paragraph(
    "X-axis変数セレクタ: Demographics全列から選択可能（目立つグラデーション枠で表示）。\n"
    "数値列→Scatter(回帰直線+R²+p値)、カテゴリ列→Boxplot に自動切替。\n"
    "Dx Boxplotは常時表示(Wilcoxon p値付き)。\n"
    "Combined scatter(全インデックスまとめ、2列グリッド)。\n"
    "Figure Settings(accordion): サイズ(1000x1000正方形), DPI(1-4x), フォント, 色, マーカー。\n"
    "PNG/SVGダウンロード。")

doc.add_heading("1.6.4 ヘルプツールチップ「?」", level=3)
doc.add_paragraph("スコア列・Boxplot・Scatterに「?」アイコン。ホバーで詳細な解釈ガイドを表示:")
bul("スコア列: ", "正/負の意味、具体例（Devで+50→胎児的、-50→成熟的）、セルクリック誘導")
bul("Boxplot: ", "箱の読み方、Controlとの比較の意味、p値の星印(***/**/*//ns)")
bul("Scatter: ", "傾きの意味、R²の解釈（1に近い=強い）、p値の意味")

doc.add_heading("1.6.5 統計検定", level=3)
add_table(["検定","用途"], [
    ["Wilcoxon(Mann-Whitney U)","Boxplot: Dx vs Control p値"],
    ["線形回帰+t検定","Scatter: R², p値"],
])

doc.add_heading("1.6.6 Methods文", level=3)
doc.add_paragraph("ボタン→論文用英文+パラメータ埋込+引用文献+クリップボードコピー。")

doc.add_heading("1.6.7 Dx色", level=3)
add_table(["Dx","色"], [["Control","#4ade80(緑)"],["Schizo","#f87171(赤)"],["Bipolar","#a78bfa(紫)"],["MDD","#fb923c(橙)"],["その他","#94a3b8(グレー)"]])

# 1.7 UI/UX
doc.add_heading("1.7 UI/UX", level=2)
doc.add_paragraph(
    "トップセクション: RFindロゴ(48px) + GitHub/How to Cite + キャッチコピー + デモ自動実行ボタン（大きく中央、青グラデーション）。\n"
    "How to Cite: 引用ボックス+コピーボタン+URL+GitHub+アルゴリズム出典。\n"
    "Accordion折りたたみ: 各Stepの入力ファイルの作り方（具体的テーブル例＋テンプレートDL）、Figure Settings。\n"
    "エラーサジェスト: Control未検出→近い値提案、ID不一致→部分一致ヒント。\n"
    "中止ボタン: 解析中に赤ボタン。\n"
    "i18n: 日/英全テキスト・エラー・ツールチップ切替。\n"
    "プライバシー: フッターで「データ外部送信なし」明記。")

# ═══ 2. 実装内容 ═══
doc.add_heading("2. 実装内容", level=1)

doc.add_heading("2.1 全機能一覧(21項目)", level=2)
for t_text, desc in [
    ("Step1: 発現量マトリックス", "CSV/TSV/Excel、検証、HGNC+Ortholog正規化、プレビュー"),
    ("Step2: サンプル情報", "ID/Dx/Age自動検出(列順不問)、ID照合、Control検証、エラーサジェスト"),
    ("Step3: 参照インデックス", "プリセット4種(出典リンク付) + カスタム(複数/D&D/Excel複数シート) + Gene Group + メタデータスキップ + FC判定"),
    ("Step4: 解析実行", "FC計算→Running Fisher(P=commonGenes)、バッチ処理、中止ボタン"),
    ("高速化", "hit-onlyスキャン。263s×4idx≈18秒"),
    ("結果テーブル", "スコア行列、ソート、色付き、「?」ツールチップ(具体例付き)"),
    ("Overlapping genes", "4方向別寄与遺伝子リスト、FC表示"),
    ("Excel出力", "Summary+インデックスごと+Parameters"),
    ("Figure(Plotly)", "X-axis変数選択(目立つUI)、Dx Boxplot(p値)、Scatter(回帰)、Combined、Settings、PNG/SVG"),
    ("Methods文", "論文用英文+パラメータ+コピー"),
    ("ツールチップ", "スコア・Boxplot・Scatter解釈ガイド(具体例、日英)"),
    ("デモデータ", "ワンクリック自動実行（「▶ デモデータで試す（数秒で結果が出ます）」）"),
    ("How to Cite", "引用ボックス+コピー+GitHub/論文リンク"),
    ("HGNC正規化", "97,737エントリ"),
    ("Ortholog", "2,042エントリ(Mouse/Rat→Human)"),
    ("テンプレート", "各StepにExcelテンプレート(テーブル例付き)"),
    ("エラーサジェスト", "Control提案、ID不一致ヒント"),
    ("Accordion", "入力ガイド/Figure Settings 折りたたみ"),
    ("i18n", "日/英全テキスト"),
    ("中止ボタン", "解析中止"),
    ("Gene Group対応", "FC列なし→全FC=1(方向性なしenrichment)"),
]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(t_text + ": "); run.bold = True; p.add_run(desc)

doc.add_heading("2.2 ファイル構成", level=2)
add_table(["ファイル","サイズ","説明"], [
    ["index.html","~100KB","メインアプリ(HTML+CSS+JS)"],
    ["hgnc_map.json","2.1MB","HGNC mapping"],
    ["ortholog_map.json","42KB","Ortholog mapping"],
    ["Dev.csv","198KB","プリセット: Dev(6,537)"],
    ["imGC_DEG.csv","15KB","プリセット: imGC(541)"],
    ["Kainate_Day03.csv","100KB","プリセット: Excitability(2,654)"],
    ["pH.csv","15KB","プリセット: pH(304)"],
    ["FPKM_demo.xlsx","8.9MB","デモ: 発現量(50サンプル)"],
    ["Demographics_demo.xlsx","16KB","デモ: サンプル情報"],
    ["scripts/build_hgnc_map.py","3KB","HGNC生成"],
    ["scripts/build_ortholog_map.py","3KB","Ortholog生成"],
])

doc.add_heading("2.3 主要関数(27関数)", level=2)
add_table(["関数","目的"], [
    ["handleFPKMFile","FPKM解析"],["handleDemoFile","サンプル情報解析"],["matchDemoIfReady","ID照合"],
    ["addIndexFromFile","インデックス(Excel複数シート)"],["parseIndexCSV","CSVパース(Gene Group対応、ヘッダー自動検出)"],
    ["detectFCFormat","FC判定"],["convertFC","FC変換"],["normalizeGene","遺伝子名正規化"],
    ["ensureHGNC","マッピング読込"],["runRFind","メイン解析"],["computeSampleFC","FC計算"],
    ["runningFisher","Running Fisher(遺伝子リスト対応)"],["phyperUpper","超幾何分布"],
    ["mannWhitneyU","Wilcoxon"],["linearRegression","線形回帰"],
    ["renderResults","テーブル(ソート)"],["drawFigures","X-axisセレクタ+描画"],
    ["drawFiguresWithVar","変数型別Figure"],["showOverlap","Overlap詳細"],
    ["dlExcel","Excel出力"],["copyMethods","Methods文"],
    ["loadDemoAndRun","デモ自動実行"],["loadPreset","プリセット読込"],
    ["cancelRFind","中止"],["applyFigSettings","Figure設定"],
    ["dlTemplate","テンプレート"],["buildPresetBtns","プリセットボタン"],
])

doc.add_heading("2.4 外部依存", level=2)
add_table(["ライブラリ","CDN","用途"], [
    ["SheetJS","cdn.sheetjs.com/xlsx-0.20.3/.../xlsx.full.min.js","Excel/CSV"],
    ["Plotly.js","cdn.jsdelivr.net/npm/plotly.js-dist-min@2","可視化"],
    ["Google Fonts","fonts.googleapis.com/css2?...","フォント"],
])

# ═══ 3. 残された課題 ═══
doc.add_heading("3. 残された課題", level=1)
doc.add_heading("3.1 優先度: 高", level=2)
add_table(["課題","詳細"], [
    ["数値検証","Rスクリプト(Chunk3_RFind.R)との数値照合。通しテスト通過済み"],
])
doc.add_heading("3.2 優先度: 中〜低", level=2)
add_table(["課題","詳細"], [
    ["プリセット拡充","CSV+PRESETS配列1行で追加"],
    ["HGNC/Ortholog更新","buildスクリプト再実行"],
    ["R/Pythonスクリプト公開","再現性のためGitHubに同梱"],
    ["ライトテーマ","CSS変数切替で対応可能"],
])

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "RFind_Web_260330_v3.docx")
out = os.path.abspath(out)
doc.save(out)
print(f"Saved: {out}")
